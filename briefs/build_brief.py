"""
DailyFruits — Creative Brief (OPEN VERSION)
Kierunek, nie ramka. Designer ma przestrzeń.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x1B, 0x5E, 0x3A)
LIME = RGBColor(0x8D, 0xC6, 0x3F)
RED = RGBColor(0xE4, 0x30, 0x20)
GRAY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

n = doc.styles['Normal']
n.font.name = 'Calibri'
n.font.size = Pt(10.5)


def H(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.bold = True
    r.font.color.rgb = GREEN
    r.font.size = Pt(15 if level == 1 else 11.5)


def P(text, bold=False, italic=False, color=BLACK, size=10.5, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.size = Pt(size)


def B(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10.5)


# ============ HEADER ============
t = doc.add_paragraph()
r = t.add_run('CREATIVE BRIEF')
r.font.bold = True
r.font.size = Pt(24)
r.font.color.rgb = GREEN

s = doc.add_paragraph()
r = s.add_run('DailyFruits  ·  Instagram social  ·  Maj 2026')
r.font.size = Pt(10)
r.font.color.rgb = GRAY
r.font.italic = True

# ============ PRODUKT ============
H('Produkt — w jednym akapicie')
P('DailyFruits dostarcza do biur owoce, kanapki, soki, finger foods, śniadania i bakalie — '
  'wszystko w jednej subskrypcji, jednej dostawie, jednej fakturze. 14 lat na rynku, '
  '2000+ klientów, w tym Google, Microsoft, ING, J.P. Morgan, Allegro. '
  'To nie catering "raz na imprezę" — to codzienny rytuał biurowy '
  '(Owocowy Czwartek, Kanapkowy Wtorek, Smoothie Monday).')

P('Co odróżnia produkt od konkurencji:', bold=True)
B('Jeden partner zamiast 5 dostawców (operacyjne odciążenie Office Managera)')
B('Subskrypcja, nie kontrakt — 0 zł kar, zmień lub anuluj kiedy chcesz')
B('Wellbeing z liczbami — pozycjonujemy się jako benefit z ROI, nie estetyczny gadget')

# ============ DLA KOGO ============
H('Dla kogo to robimy')
P('Decydenci B2B w firmach 50–500 osób. Trzy główne typy odbiorców:')
B('HR / People & Culture — szuka wellbeing benefitu, który da się obronić liczbami przed zarządem')
B('Office Manager — operacyjnie ogarnia codzienność biura, ma dość 5 dostawców i 5 faktur')
B('CEO / Founder — pyta o ROI i status (czy robią to inne duże marki)')

P('Każdy z nich reaguje inaczej. Możesz wybrać, do kogo celuje który materiał — '
  'lub zrobić serię, która działa na całym lejku.', italic=True, color=GRAY)

# ============ INSIGHT ============
H('Insight, na którym budujemy')
P('Głód w biurze to operacyjny koszt, którego nikt nie mierzy. 68% pracowników jest mniej '
  'efektywnych, gdy odczuwa głód. 50% nie je zdrowo. 57% chce, by pracodawca w tym pomógł. '
  'Większość firm zostawia to przypadkowi — automatom z batonami, kawie z saszetek, lunchowi '
  'na mieście. DailyFruits zamienia to w przewidywalny, codzienny rytuał.')

P('Tak naprawdę sprzedajemy: spokój operacyjny, status pracodawcy, mierzalny benefit. '
  'A nie owoce.', italic=True, color=RED)

# ============ TON & ESTETYKA ============
H('Ton i estetyka — kierunek, nie ramka')
P('Marka jest jasna, ciepła, lekko nieformalna, ale mówi językiem biznesu — '
  'liczby, konkret, krótkie zdania. Bez "smacznych chwil dla Twojego zespołu". '
  'Bez stocków uśmiechniętych ludzi w open space. Bez wellnessowej papki.')

P('Z grubsza paleta:', bold=True)
B('Zieleń ciemna (#1B5E3A), lime (#8DC63F), cream (#FFF9E0) — baza')
B('Yellow (#FFF200), red (#E43020) — akcent / CTA')

P('Z grubsza typografia:', bold=True)
B('DM Sans — body')
B('Achiko / Lobster — fun-fontowe headline i duże liczby (kotwica marki)')

P('Pomocne assety:', bold=True)
B('Hand-painted ilustracje (MJ-style) — folder /icons/ na stronie, 30 sztuk')
B('Packshoty produktowe — folder /packshot-*.webp')
B('Logotypy klientów — sekcja "Zaufali nam liderzy rynku" na dailyfruits.pl')

P('Sygnatura wizualna marki: stickery (obrócone chipy z fun-font napisem) i organiczne fale '
  'między sekcjami. Możesz to wykorzystać lub zaproponować coś, co rozwija ten DNA.', italic=True, color=GRAY)

# ============ CO DOSTARCZASZ ============
doc.add_page_break()
H('Co dostarczasz')

H('A. 6 postów na feed Instagram', level=2)
P('Spójna seria, która razem opowiada historię produktu — od problemu, przez rozwiązanie, '
  'do działania. Format zostawiamy Ci do wyboru (1:1 lub 4:5 — co ma większy sens dla danego posta).')

P('Kierunki tematyczne (możesz je przemieszać, połączyć, rozwinąć po swojemu):', bold=True)
B('Problem zespołu — głód, sugar crash, spadek koncentracji po 14:00')
B('Social proof — wielcy klienci, opinie, liczba lat na rynku')
B('Operacyjne odciążenie — jeden partner vs. chaos 5 dostawców')
B('Rytuał biurowy — Owocowy Czwartek, Kanapkowy Wtorek, Smoothie Monday')
B('Elastyczność — subskrypcja, nie kontrakt, 0 zł kar')
B('Wezwanie do działania — kalkulator dailyfruits.pl/kalkulator (30 sekund, bez handlowca)')

P('Nie musisz robić 1:1 mapowania "6 kierunków = 6 postów". Możesz złączyć dwa, rozbić jeden na dwa, '
  'znaleźć siódmy kąt, który lepiej działa. Chcemy efektu, nie listy do odhaczenia.', italic=True, color=GRAY)

H('B. 1 karuzela edukacyjna (Instagram)', level=2)
P('Format 4:5, długość: tyle slajdów ile potrzeba (6–10). Cel: save-worthy content, który ktoś '
  'z HR wysyła dalej w firmie.')

P('Temat: dlaczego warto, by zespół jadł w biurze zdrowo — z liczbami, bez wellnessowej papki. '
  'Możesz to opowiedzieć jako "X powodów", jako mini-case study, jako wywód oparty na stacie. '
  'Wybierz formę, która najmocniej pociągnie scroll.')

P('Pomocne liczby (użyj tych, które najlepiej grają w Twojej historii):', bold=True)
B('68% — pracownicy mniej efektywni przy głodzie')
B('50% — nie je zdrowo')
B('57% — chce wsparcia pracodawcy w zdrowym jedzeniu')
B('36% — pracuje regularnie z uczuciem głodu')
B('55% — próbuje zmienić nawyki na zdrowsze')
B('2000+ firm, 14 lat, ★ 4.9/5 (120+ opinii)')

H('C. 2 linie kampanijne (paid)', level=2)
P('Dwie różne historie sprzedażowe — każda spójna seria w 3 formatach (Story/Reels 9:16, '
  'Feed kwadrat 1:1, LinkedIn 1200×628 lub 1200×1200).')

P('Linia 1 — Performance / Direct Response:', bold=True)
P('Cel: kliki w kalkulator, leady. Mocny hook, twarda liczba, problem-aware audience. '
  'Brzmienie: konkret, taktyka, lekka prowokacja. Można iść z perspektywy "głód kosztuje firmę", '
  '"zespół pracuje gorzej, niż myślisz" albo z innego kąta, który Ci się otworzy.', size=10)

P('Linia 2 — Brand / Aspiracyjna:', bold=True)
P('Cel: zasięg, zapamiętywalność, zaufanie. Cold audience — pierwsza styczność z marką. '
  'Brzmienie: status, social proof, premium, lekka klasa. Można iść z perspektywy "dołącz do klubu '
  'Microsoft/Google/ING", "rytuał, który robi się od 14 lat" albo z czegoś, co Ci sprawi większą '
  'frajdę.', size=10)

P('Linie powinny być na tyle różne, by ktoś, kto zobaczy obie reklamy, nie pomyślał "to ten sam komunikat". '
  'Inny ton, inna estetyka, inny rytm.', italic=True, color=GRAY)

# ============ ZASADY ============
H('Kilka zasad, których się trzymamy')
B('Każdy materiał ma minimum 1 element w fun-fonzie (Achiko/Lobster) — to kotwica marki')
B('Liczby tak — przymiotniki na siłę nie ("pyszne", "soczyste", "smaczne chwile" → out)')
B('Max 1–2 emotki na post')
B('Jasna, ciepła paleta. Nie schodzimy w dark mode, nie idziemy w gradienty/neony')
B('CTA prowadzi do dailyfruits.pl lub konkretnej podstrony (oferta, kalkulator, dostawa)')

# ============ LOGISTYKA ============
H('Logistyka')

P('Eksport:', bold=True)
B('PNG (q80) + JPG (q90) per format')
B('Master Figma lub Illustrator z osadzonymi fontami')
B('Nazewnictwo: DF_2026-06_NAZWA_FORMAT_v01.png')

P('Timeline:', bold=True)
B('Faza koncepcji — 3 dni (3 kierunki na linię kampanijną, moodboard, sketches postów)')
B('Akceptacja kierunku — 1 dzień')
B('Produkcja — 4 dni')
B('Poprawki — 2 rundy, 2 dni')
B('Łącznie: ok. 10 dni roboczych od kick-offu')

P('Brand assety do pobrania:', bold=True)
B('dailyfruits.pl — produkcyjny system wizualny, kolory w shared.css')
B('Folder /icons/ — 30 hand-painted webp')
B('Folder /packshot-*.webp — produkty na białym tle')

# Footer
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
r = p.add_run('Pytania → Reszek · p.reszkovy@gmail.com')
r.font.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GRAY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
import os
output = '/sessions/zen-modest-ride/mnt/Fruityyyy/briefs/DailyFruits-Creative-Brief.docx'
os.makedirs(os.path.dirname(output), exist_ok=True)
doc.save(output)
print(f'OK: {output} ({os.path.getsize(output)} bytes)')
